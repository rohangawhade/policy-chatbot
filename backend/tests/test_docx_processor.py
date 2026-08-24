from pathlib import Path

import pytest
from docx import Document as DocxDocument

from adapters.document_processors.docx_processor import DOCXProcessor
from core.ports.document_processor_port import DocumentProcessorPort


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.core_properties.title = "Sample SPD"
    document.core_properties.author = "PolicyPal"
    document.add_paragraph("Summary Plan Description")
    document.add_paragraph("Your annual deductible is $500.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Plan"
    table.rows[0].cells[1].text = "Dental PPO"
    document.save(path)
    return path


def test_is_a_document_processor_port() -> None:
    assert isinstance(DOCXProcessor(), DocumentProcessorPort)


def test_extract_text_returns_paragraph_and_table_content(sample_docx: Path) -> None:
    text = DOCXProcessor().extract_text(str(sample_docx))

    assert "Summary Plan Description" in text
    assert "Your annual deductible is $500." in text
    assert "Plan | Dental PPO" in text


def test_extract_text_skips_empty_paragraphs(tmp_path: Path) -> None:
    path = tmp_path / "with_blanks.docx"
    document = DocxDocument()
    document.add_paragraph("First line")
    document.add_paragraph("")
    document.add_paragraph("Second line")
    document.save(path)

    text = DOCXProcessor().extract_text(str(path))

    assert text == "First line\nSecond line"


def test_extract_metadata_returns_counts_and_title(sample_docx: Path) -> None:
    metadata = DOCXProcessor().extract_metadata(str(sample_docx))

    assert metadata["paragraph_count"] == 2
    assert metadata["table_count"] == 1
    assert metadata["title"] == "Sample SPD"
    assert metadata["author"] == "PolicyPal"


def test_extract_metadata_returns_none_for_missing_title(tmp_path: Path) -> None:
    path = tmp_path / "untitled.docx"
    DocxDocument().save(path)

    metadata = DOCXProcessor().extract_metadata(str(path))

    assert metadata["title"] is None
