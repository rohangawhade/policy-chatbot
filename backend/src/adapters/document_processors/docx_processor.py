"""DOCX document processor: python-docx for paragraph and table text
extraction (files/plan.md Step 3.6)."""

from typing import Any

from docx import Document as DocxDocument

from core.ports.document_processor_port import DocumentProcessorPort


class DOCXProcessor(DocumentProcessorPort):
    def extract_text(self, file_path: str) -> str:
        document = DocxDocument(file_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = " | ".join(cell.text.strip() for cell in row.cells)
                if cells.strip(" |"):
                    parts.append(cells)
        return "\n".join(parts)

    def extract_metadata(self, file_path: str) -> dict[str, Any]:
        document = DocxDocument(file_path)
        core_properties = document.core_properties
        return {
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "title": core_properties.title or None,
            "author": core_properties.author or None,
        }
