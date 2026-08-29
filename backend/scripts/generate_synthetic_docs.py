"""Generates synthetic employer benefits documents using an LLM (files/plan.md
Step 11.2), for the same 5 fictional employers `scripts/seed_data.py` seeds
(Step 11.3) -- `seed_data.py`'s own upload step recursively discovers
everything under `data/synthetic/` alongside `data/gov_pdfs/` (Step 11.1),
so these give each demo employer real, on-topic content beyond the
government-generic PDFs.

10 document types per employer (50 total, matching this step's "50+" target):
one plan-summary document per `PolicyType` (health/dental/vision/life/
disability -- covering exactly what the RAG pipeline's `policy_type` metadata
filter and `_detect_policy_type()` keyword matching need to be meaningfully
exercised) plus five employer-wide documents (handbook benefits section, open
enrollment guide, benefits FAQ, COBRA continuation guide, wellness program
guide) that plan.md's own bullet list names or clearly implies. Formats
alternate DOCX/PDF across the 10 types (5 each) -- exercising both real
document processors (Step 3.6) with genuinely LLM-authored content, not just
the four hand-curated real government documents Step 11.1 already covers.

Section headings in the generated content are deliberately Title Case with no
trailing punctuation, under `_HEADING_MAX_LENGTH` -- `MetadataExtractor`
(Step 4.1) detects headings by exactly this heuristic (checked against its
source directly, not assumed), so structuring generation output to match it
gives these documents the same `section_title` chunk metadata a real,
well-formatted document would get, rather than every chunk landing with
`section_title=None`.

Idempotent by design, same manifest pattern as `download_gov_docs.py` (Step
11.1): `data/synthetic/manifest.json` records every file already generated;
re-running only fills in what's missing unless --force is passed. Each call
is a real, billed LLM request (however cheap) -- idempotency here also avoids
silently regenerating (and re-paying for) content that already exists.

Usage:
    python scripts/generate_synthetic_docs.py [--dry-run] [--force]
        [--employer SLUG] [--doc-type KEY] [--limit N]

    --dry-run     List what would be generated without calling the LLM.
    --force       Regenerate even documents the manifest already has.
    --employer    Restrict to one employer (slug, e.g. "northwind-traders").
    --doc-type    Restrict to one document type (see _DOC_SPECS keys below).
    --limit N     Cap the number of documents generated this run (testing).
"""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import io
import json
import re
import sys
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Literal

import fitz
import structlog
from docx import Document as DocxDocument

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from adapters.llm.litellm_adapter import LiteLLMAdapter  # noqa: E402
from config import llm_config  # noqa: E402
from core.domain.policy import PolicyType  # noqa: E402

logger = structlog.get_logger(__name__)

_REPO_ROOT = Path(__file__).resolve().parents[2]
_OUTPUT_ROOT = _REPO_ROOT / "data" / "synthetic"
_MANIFEST_PATH = _OUTPUT_ROOT / "manifest.json"

# Must match scripts/seed_data.py's _COMPANY_NAMES exactly -- these two
# scripts have no other dependency on each other, and a 5-line constant
# isn't worth a shared module, but the names need to line up so seed_data.py
# uploads this content for the same employers it creates.
_COMPANY_NAMES = [
    "Northwind Traders",
    "Globex Corporation",
    "Acme Manufacturing",
    "Initech Solutions",
    "Contoso Health Group",
]

_HEADING_MAX_LENGTH = 80  # Must match adapters/chunking/metadata_extractor.py
_GENERATION_TEMPERATURE = 0.7  # Higher than LLMConfig's 0.1 default -- these
# are meant to read as 50 distinct real documents, not near-identical ones.
_MAX_TOKENS = 1500

# Groq's free tier caps `openai/gpt-oss-20b` at 8000 tokens/minute (TPM), and
# each request here uses ~1800 (prompt + completion) -- empirically confirmed
# via a real run that failed 26/50 documents with GroqException rate_limit_
# exceeded once TPM usage climbed past ~7000. A fixed gap between requests
# keeps sustained usage under that ceiling instead of relying on retries
# (LiteLLMAdapter's built-in retry/backoff isn't long enough for an 8s+ TPM
# reset window on every one of 50 sequential calls).
_REQUEST_INTERVAL_SECONDS = 15.0


def _slugify(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")


@dataclass(frozen=True, kw_only=True)
class DocSpec:
    """One document type this script generates once per employer."""

    key: str
    title_template: str
    covers: str
    policy_type: PolicyType | None
    format: Literal["docx", "pdf"]


_DOC_SPECS: list[DocSpec] = [
    DocSpec(
        key="health_plan_summary",
        title_template="{company} Health Plan Summary",
        covers=(
            "a summary of the company's group health insurance plan: coverage tiers "
            "(employee only, employee+spouse, family), monthly premiums, annual "
            "deductible, out-of-pocket maximum, copay amounts for primary care/specialist/"
            "urgent care/ER visits, and what's covered vs. excluded"
        ),
        policy_type=PolicyType.HEALTH,
        format="docx",
    ),
    DocSpec(
        key="dental_plan_summary",
        title_template="{company} Dental Plan Summary",
        covers=(
            "a summary of the company's dental insurance plan: preventive care coverage "
            "(cleanings, exams, X-rays), basic procedures (fillings, extractions), major "
            "procedures (crowns, root canals, dentures), the annual maximum benefit, and "
            "any waiting periods"
        ),
        policy_type=PolicyType.DENTAL,
        format="pdf",
    ),
    DocSpec(
        key="vision_plan_summary",
        title_template="{company} Vision Plan Summary",
        covers=(
            "a summary of the company's vision insurance plan: annual eye exam coverage, "
            "frame and lens allowances, contact lens allowance, discounts on LASIK, and "
            "in-network vs out-of-network provider differences"
        ),
        policy_type=PolicyType.VISION,
        format="docx",
    ),
    DocSpec(
        key="life_insurance_summary",
        title_template="{company} Life Insurance Plan Summary",
        covers=(
            "a summary of the company's group life insurance plan: the basic "
            "employer-paid coverage amount, voluntary supplemental coverage options and "
            "their cost, accidental death and dismemberment (AD&D) coverage, and how to "
            "name or update a beneficiary"
        ),
        policy_type=PolicyType.LIFE,
        format="pdf",
    ),
    DocSpec(
        key="disability_insurance_summary",
        title_template="{company} Disability Insurance Plan Summary",
        covers=(
            "a summary of the company's disability insurance plans: short-term "
            "disability (benefit percentage, maximum duration, elimination period) and "
            "long-term disability (benefit percentage, maximum duration, definition of "
            "disability used), and how to file a claim"
        ),
        policy_type=PolicyType.DISABILITY,
        format="docx",
    ),
    DocSpec(
        key="employee_handbook_benefits",
        title_template="{company} Employee Handbook - Benefits Section",
        covers=(
            "the benefits chapter of an employee handbook: an overview of all benefit "
            "programs offered, eligibility requirements (e.g. full-time status, waiting "
            "period after hire), how and when benefits can be changed outside open "
            "enrollment (qualifying life events), and who to contact with questions"
        ),
        policy_type=None,
        format="pdf",
    ),
    DocSpec(
        key="open_enrollment_guide",
        title_template="{company} Open Enrollment Guide",
        covers=(
            "an open enrollment guide: the enrollment window dates, step-by-step "
            "instructions for enrolling or changing elections online, what happens if an "
            "employee takes no action, and a comparison of this year's plan changes "
            "versus last year"
        ),
        policy_type=None,
        format="docx",
    ),
    DocSpec(
        key="benefits_faq",
        title_template="{company} Benefits FAQ",
        covers=(
            "a frequently-asked-questions document about employee benefits, covering at "
            "least 8 distinct questions employees commonly ask about health, dental, "
            "vision, and other benefits (for example: when coverage starts, adding a "
            "new spouse or dependent mid-year, the difference between an HMO and a PPO), "
            "each with a clear, direct answer"
        ),
        policy_type=None,
        format="pdf",
    ),
    DocSpec(
        key="cobra_continuation_guide",
        title_template="{company} COBRA Continuation Coverage Guide",
        covers=(
            "a guide explaining COBRA continuation coverage: who is eligible after a "
            "qualifying event (termination, reduced hours, divorce, etc.), how long "
            "coverage can continue, how premiums are calculated, and the deadline to "
            "elect COBRA after losing coverage"
        ),
        policy_type=None,
        format="docx",
    ),
    DocSpec(
        key="wellness_program_guide",
        title_template="{company} Wellness Program Guide",
        covers=(
            "a guide to the company's employee wellness program: the programs available "
            "(for example gym reimbursement, biometric screenings, an employee assistance "
            "program), any premium discounts or rewards for participating, and how to "
            "enroll in each program"
        ),
        policy_type=None,
        format="pdf",
    ),
]

_DOC_SPECS_BY_KEY = {spec.key: spec for spec in _DOC_SPECS}


@dataclass(frozen=True)
class GeneratedSection:
    heading: str
    body: str


def _build_prompt(company: str, spec: DocSpec) -> str:
    return (
        f"Write {spec.covers} for {company}, a fictional company. This is test data for "
        "a software demo -- invent specific, realistic numbers (dollar amounts, "
        "percentages, dates) and never use placeholders like '[amount]' or 'TBD'. Do not "
        "mention anywhere that this is fictional, a demo, or AI-generated -- write it "
        "exactly as a real HR department would.\n\n"
        "Format your response EXACTLY like this, with no text before the first heading "
        "and no text after the last section:\n\n"
        "## Section Heading One\n"
        "Body paragraph text for this section, 2-4 sentences.\n\n"
        "## Section Heading Two\n"
        "Body paragraph text for this section, 2-4 sentences.\n\n"
        "Use 4 to 7 sections. Each section heading must be Title Case, under 8 words, "
        f"and no more than {_HEADING_MAX_LENGTH} characters, with no trailing punctuation."
    )


# Real, empirically-found rendering bug, not a hypothetical: PyMuPDF's
# base-14 "helv" font (used by _render_pdf below) has no glyphs for these --
# each one renders as a literal "?" in the actual PDF page content (verified
# with a minimal repro: insert_textbox() + extract_text() round-tripped
# "X-rays" back as "X?rays"). LLM output uses these constantly, so every
# generated document needs this before rendering -- applied to both PDF and
# DOCX text for one consistent look, not just the format that's actually
# broken without it.
_UNICODE_PUNCTUATION_TO_ASCII = {
    "‘": "'",
    "’": "'",
    "“": '"',
    "”": '"',
    "–": "-",
    "—": "-",
    "‑": "-",
    "…": "...",
    " ": " ",  # narrow no-break space -- e.g. the model's "100 %"
    " ": " ",  # non-breaking space
    " ": " ",  # thin space
}


def _to_ascii_punctuation(text: str) -> str:
    for unicode_char, ascii_char in _UNICODE_PUNCTUATION_TO_ASCII.items():
        text = text.replace(unicode_char, ascii_char)
    return text


def _parse_sections(raw_text: str) -> list[GeneratedSection]:
    sections: list[GeneratedSection] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        body = "\n".join(current_lines).strip()
        if current_heading and body:
            sections.append(GeneratedSection(current_heading, body))

    for line in _to_ascii_punctuation(raw_text).splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            flush()
            current_heading = _normalize_heading(stripped.removeprefix("## ").strip())
            current_lines = []
        elif stripped:
            current_lines.append(stripped)
    flush()

    if not sections:
        # Defensive fallback -- the model didn't follow the "## heading"
        # format. Real LLM output, not a hypothetical: worth handling
        # rather than failing this document outright.
        logger.warning("model_did_not_follow_section_format")
        sections = [GeneratedSection("Overview", raw_text.strip())]
    return sections


def _normalize_heading(heading: str) -> str:
    # MetadataExtractor (Step 4.1) rejects a heading candidate that ends in
    # trailing punctuation or exceeds _HEADING_MAX_LENGTH -- strip/truncate
    # defensively rather than trust the model followed the prompt exactly.
    stripped = heading.rstrip(".,;:")
    return stripped[:_HEADING_MAX_LENGTH]


def _render_docx(title: str, sections: list[GeneratedSection]) -> bytes:
    document = DocxDocument()
    document.add_heading(title, level=1)
    for section in sections:
        document.add_heading(section.heading, level=2)
        for paragraph in section.body.split("\n"):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_PDF_PAGE_SIZE = (612, 792)  # US Letter, points
_PDF_MARGIN = 50
_PDF_RECT = fitz.Rect(
    _PDF_MARGIN, _PDF_MARGIN, _PDF_PAGE_SIZE[0] - _PDF_MARGIN, _PDF_PAGE_SIZE[1] - _PDF_MARGIN
)


def _render_pdf(title: str, sections: list[GeneratedSection]) -> bytes:
    # One section per page -- sidesteps manual pagination math entirely
    # (`insert_textbox` draws what fits and silently discards the rest; it
    # doesn't hand back leftover text to flow onto a new page). Sections are
    # prompted to stay short (2-4 sentences) specifically so this fits
    # within one page's rect at a normal font size.
    document = fitz.open()
    for index, section in enumerate(sections):
        page = document.new_page(width=_PDF_PAGE_SIZE[0], height=_PDF_PAGE_SIZE[1])
        heading_prefix = f"{title}\n\n" if index == 0 else ""
        content = f"{heading_prefix}{section.heading}\n\n{section.body}"
        overflow = page.insert_textbox(_PDF_RECT, content, fontsize=11, fontname="helv")
        if overflow < 0:
            logger.warning(
                "pdf_section_overflow",
                title=title,
                heading=section.heading,
                hint="Section content was longer than one page -- truncated.",
            )
    result = document.tobytes()
    if not isinstance(result, bytes):
        raise TypeError(f"expected bytes from fitz Document.tobytes(), got {type(result).__name__}")
    return result


@dataclass
class Manifest:
    entries: dict[str, dict[str, str | int]] = field(default_factory=dict)

    @classmethod
    def load(cls) -> Manifest:
        if not _MANIFEST_PATH.exists():
            return cls()
        return cls(entries=json.loads(_MANIFEST_PATH.read_text(encoding="utf-8")))

    def save(self) -> None:
        _MANIFEST_PATH.parent.mkdir(parents=True, exist_ok=True)
        _MANIFEST_PATH.write_text(json.dumps(self.entries, indent=2, sort_keys=True) + "\n")

    def has(self, relative_path: str) -> bool:
        return relative_path in self.entries

    def record(self, relative_path: str, *, model: str, content: bytes) -> None:
        self.entries[relative_path] = {
            "model": model,
            "sha256": hashlib.sha256(content).hexdigest(),
            "size_bytes": len(content),
            "generated_at": datetime.now(UTC).isoformat(),
        }


async def _generate_one(llm: LiteLLMAdapter, company: str, spec: DocSpec) -> bytes:
    title = spec.title_template.format(company=company)
    prompt = _build_prompt(company, spec)
    raw_text = await llm.generate(
        prompt,
        model=llm_config.cheap_model,
        temperature=_GENERATION_TEMPERATURE,
        max_tokens=_MAX_TOKENS,
    )
    sections = _parse_sections(raw_text)
    if spec.format == "docx":
        return _render_docx(title, sections)
    return _render_pdf(title, sections)


async def _generate_all(
    *,
    dry_run: bool,
    force: bool,
    employer_filter: str | None,
    doc_type_filter: str | None,
    limit: int | None,
) -> tuple[int, int, int]:
    """Returns (generated, skipped_existing, failed)."""
    manifest = Manifest.load()
    llm = LiteLLMAdapter()
    generated = skipped = failed = 0
    attempted = 0

    for company in _COMPANY_NAMES:
        slug = _slugify(company)
        if employer_filter and slug != employer_filter:
            continue

        for spec in _DOC_SPECS:
            if doc_type_filter and spec.key != doc_type_filter:
                continue
            if limit is not None and attempted >= limit:
                return generated, skipped, failed

            destination = _OUTPUT_ROOT / slug / f"{spec.key}.{spec.format}"
            relative_path = destination.relative_to(_OUTPUT_ROOT).as_posix()

            if not force and destination.exists() and manifest.has(relative_path):
                skipped += 1
                continue

            attempted += 1
            if dry_run:
                logger.info("would_generate", path=relative_path)
                generated += 1
                continue

            if attempted > 1:
                await asyncio.sleep(_REQUEST_INTERVAL_SECONDS)

            try:
                content = await _generate_one(llm, company, spec)
            except Exception as exc:
                logger.warning("generation_failed", path=relative_path, error=str(exc))
                failed += 1
                continue

            destination.parent.mkdir(parents=True, exist_ok=True)
            destination.write_bytes(content)
            manifest.record(relative_path, model=llm_config.cheap_model, content=content)
            manifest.save()
            generated += 1
            logger.info("generated", path=relative_path, bytes=len(content))

    return generated, skipped, failed


async def _main(args: argparse.Namespace) -> int:
    generated, skipped, failed = await _generate_all(
        dry_run=args.dry_run,
        force=args.force,
        employer_filter=args.employer,
        doc_type_filter=args.doc_type,
        limit=args.limit,
    )
    logger.info("done", generated=generated, skipped=skipped, failed=failed)
    return 1 if failed and generated == 0 else 0


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--dry-run", action="store_true", help="List targets without calling the LLM."
    )
    parser.add_argument(
        "--force", action="store_true", help="Regenerate documents already in the manifest."
    )
    parser.add_argument(
        "--employer",
        default=None,
        help="Restrict to one employer slug (e.g. 'acme-manufacturing').",
    )
    parser.add_argument(
        "--doc-type",
        default=None,
        choices=sorted(_DOC_SPECS_BY_KEY),
        help="Restrict to one document type.",
    )
    parser.add_argument("--limit", type=int, default=None, help="Cap documents generated this run.")
    return parser.parse_args()


if __name__ == "__main__":
    sys.exit(asyncio.run(_main(_parse_args())))
